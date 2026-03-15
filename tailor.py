"""
Resume tailoring pipeline.
1. Drop a new JD PDF into new_jobs/ (or pass --jd <path>).
2. Finds the top 3 most similar past JDs in ChromaDB.
3. Locates the DOCX resume used for the closest match.
4. Calls Claude to rewrite the full resume tailored for the new job.
5. Saves the tailored resume as a DOCX in new_jobs/.

Quickest usage — drop a PDF into new_jobs/ then run:
  conda run -n job-rag python tailor.py

Or point directly at a file:
  conda run -n job-rag python tailor.py --jd path/to/new_jd.pdf
"""
import argparse
import csv
import json
import os
import re
import sys

import anthropic
import chromadb
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PyPDF2 import PdfReader

import config


# ── helpers ──────────────────────────────────────────────────────────────────

def read_pdf(path: str, max_chars: int = 8000) -> str:
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        text += (page.extract_text() or "") + " "
        if len(text) >= max_chars:
            break
    return text[:max_chars].strip()


def read_text_file(path: str, max_chars: int = 8000) -> str:
    with open(path, encoding="utf-8") as f:
        return f.read(max_chars)


def read_docx_structured(path: str) -> list[dict]:
    """Return list of {style, text} dicts preserving paragraph structure."""
    doc = Document(path)
    return [
        {"style": p.style.name, "text": p.text}
        for p in doc.paragraphs
        if p.text.strip()
    ]


def read_docx_text(path: str, max_chars: int = 8000) -> str:
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return text[:max_chars]


def extract_jd_text(path: str) -> str:
    """Extract text from a JD PDF or TXT file."""
    if not os.path.exists(path):
        sys.exit(f"File not found: {path}")
    if path.lower().endswith(".pdf"):
        text = read_pdf(path)
    else:
        text = read_text_file(path)
    if not text:
        sys.exit(f"Could not extract text from: {path}")
    return text


def collect_inbox_pdfs() -> list[str]:
    """Return all PDF paths in new_jobs/, oldest-first so output order is intuitive."""
    inbox = config.NEW_JD_DIR
    os.makedirs(inbox, exist_ok=True)
    return sorted(
        [os.path.join(inbox, f) for f in os.listdir(inbox) if f.lower().endswith(".pdf")],
        key=os.path.getmtime,
    )


def find_resume_for_match(job_folder: str, matched_path: str) -> tuple[str, str] | tuple[None, None]:
    """
    Find the resume for a matched JD. Prefers DOCX over PDF.
    Searches resume_reference.csv for likely_type=resume in the same job_folder.
    Returns (resume_path, resume_filename) or (None, None).
    """
    ref_csv = config.RESUME_REFERENCE_PATH
    if not os.path.exists(ref_csv):
        return None, None

    leaf = job_folder.split("/")[-1] if "/" in job_folder else job_folder

    docx_match = None
    pdf_match = None
    with open(ref_csv, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            if row.get("likely_type") != "resume":
                continue
            if leaf and leaf not in row["path"]:
                continue
            if row.get("extension") == "docx" and docx_match is None:
                docx_match = (row["path"], row["filename"])
            elif row.get("extension") == "pdf" and pdf_match is None:
                pdf_match = (row["path"], row["filename"])

    if docx_match:
        return docx_match
    if pdf_match:
        return pdf_match

    # Fallback: same directory as the matched JD
    jd_dir = os.path.dirname(matched_path)
    with open(ref_csv, newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            if row.get("likely_type") != "resume":
                continue
            if os.path.dirname(row["path"]) == jd_dir:
                return row["path"], row["filename"]

    return None, None


def call_claude_rewrite(new_jd: str, matched_jd: str, resume_paragraphs: list[dict]) -> list[dict]:
    """
    Ask Claude to rewrite the full resume tailored for the new JD.
    Returns a list of {style, text} paragraph dicts.
    """
    client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)

    system = (
        "You are an expert resume writer. Rewrite the given resume to be perfectly tailored "
        "for a new job description.\n\n"
        "STRICT LENGTH RULE: The final resume MUST fit on 2 pages. To achieve this:\n"
        "  - Write a concise 2-3 sentence professional summary.\n"
        "  - Include only the 3-4 most recent or most relevant roles.\n"
        "  - Max 6 bullet points per role. Each bullet max 50 words.\n"
        "  - Keep the Skills section to 1-2 compact lines.\n"
        "  - Omit roles older than 15 years unless uniquely relevant.\n\n"
        "BOLD RULE: Use **text** markdown bold markers for emphasis in Normal and List Bullet "
        "paragraphs only — e.g. key metrics (**30% faster**), critical skills (**Python**, **SQL**), "
        "and company names in Normal date/context lines. Do NOT use ** inside Heading paragraphs.\n\n"
        "OUTPUT FORMAT: Return ONLY a valid JSON array of paragraph objects with 'style' and 'text'.\n"
        "Allowed style values:\n"
        "  'Heading 1'   — candidate name (first para) and ALL-CAPS section headers. "
        "Section order MUST be: name, contact info, SUMMARY, SKILLS, EXPERIENCE, EDUCATION\n"
        "  'Heading 2'   — job title line: 'Title | Company | Start–End'\n"
        "  'Normal'      — contact info, inline context, date ranges\n"
        "  'List Bullet' — achievement bullet points\n"
        "Preserve all factual details (dates, company names, degrees, numbers). "
        "No text outside the JSON array."
    )

    resume_tagged = "\n".join(
        f"[{p['style']}] {p['text']}" for p in resume_paragraphs
    )

    user = f"""## Source Resume
{resume_tagged}

## Past Job Description (what this resume was written for)
{matched_jd}

## New Job Description (tailor the resume for this)
{new_jd}

Rewrite the full resume as a compact JSON array of paragraph objects. Remember: strict 2-page limit."""

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4096,
        messages=[{"role": "user", "content": user}],
        system=system,
    )
    raw = message.content[0].text.strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = re.search(r"\[.*\]", raw, re.DOTALL)
        if m:
            return json.loads(m.group())
        raise ValueError(f"Claude returned unexpected output:\n{raw[:500]}")


def _add_runs_with_bold(para, text: str) -> None:
    """Add runs to a paragraph, converting **bold** markers into real bold runs."""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            para.add_run(part[2:-2]).bold = True
        elif part:
            para.add_run(part)


def write_tailored_docx(paragraphs: list[dict], output_path: str) -> None:
    """Write a clean, well-structured DOCX easy to copy-paste into a template."""
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin    = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin   = Inches(1.0)
    sec.right_margin  = Inches(1.0)

    # Normal — contact lines, date ranges
    ns = doc.styles['Normal']
    ns.font.name = 'Calibri'
    ns.font.size = Pt(10.5)
    ns.paragraph_format.space_before = Pt(0)
    ns.paragraph_format.space_after  = Pt(2)

    # Heading 1 — name + ALL-CAPS section headers
    h1 = doc.styles['Heading 1']
    h1.font.name      = 'Calibri'
    h1.font.size      = Pt(14)
    h1.font.bold      = True
    h1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # professional dark blue
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after  = Pt(2)
    h1.paragraph_format.keep_with_next = True

    # Heading 2 — job title | company | dates
    h2 = doc.styles['Heading 2']
    h2.font.name      = 'Calibri'
    h2.font.size      = Pt(10.5)
    h2.font.bold      = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after  = Pt(1)
    h2.paragraph_format.keep_with_next = True

    # List Bullet — achievement bullets
    lb = doc.styles['List Bullet']
    lb.font.name = 'Calibri'
    lb.font.size = Pt(10.5)
    lb.paragraph_format.space_before   = Pt(0)
    lb.paragraph_format.space_after    = Pt(2)
    lb.paragraph_format.left_indent    = Inches(0.25)

    # Remove the default blank paragraph Word inserts
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    VALID = {'Heading 1', 'Heading 2', 'Normal', 'List Bullet'}
    for item in paragraphs:
        style = item.get('style', 'Normal')
        text  = item.get('text', '').strip()
        if not text:
            continue
        if style not in VALID:
            style = 'Normal'
        para = doc.add_paragraph(style=style)
        _add_runs_with_bold(para, text)
        # Horizontal rule under Heading 1 section headers (not the name)
        if style == 'Heading 1' and text.isupper():
            from docx.oxml.ns import qn
            from docx.oxml   import OxmlElement
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '1F497D')
            pBdr.append(bottom)
            pPr.append(pBdr)

    doc.save(output_path)


# ── main ─────────────────────────────────────────────────────────────────────

def process_one_jd(jd_path: str, collection, top_k: int) -> None:
    """Run the full tailor pipeline for a single JD PDF."""
    jd_stem = os.path.splitext(os.path.basename(jd_path))[0]
    print(f"\nLoading: {os.path.basename(jd_path)}")
    new_jd_text = extract_jd_text(jd_path)
    print(f"  {len(new_jd_text)} characters.")

    # Query ChromaDB
    results = collection.query(
        query_texts=[new_jd_text],
        n_results=min(top_k, collection.count()),
        include=["metadatas", "distances", "documents"],
    )
    metadatas = results["metadatas"][0]
    distances = results["distances"][0]
    documents = results["documents"][0]

    print(f"\nTop {len(metadatas)} similar past jobs:")
    for i, (meta, dist) in enumerate(zip(metadatas, distances), 1):
        similarity = 1 - dist
        print(f"  {i}. [{similarity:.3f}] {meta['filename']}")
        print(f"       folder: {meta['job_folder']}  outcome: {meta['outcome']}")

    top_meta = metadatas[0]
    top_doc  = documents[0]
    print(f"\nUsing top match: {top_meta['filename']}")

    # Find source resume (prefer DOCX)
    resume_path, resume_filename = find_resume_for_match(
        top_meta["job_folder"], top_meta["path"]
    )
    if not resume_path:
        print("  [skip] No matching resume found — skipping this JD.")
        return

    ext = os.path.splitext(resume_filename)[1].lower()
    print(f"Source resume: {resume_filename}  ({ext})")

    try:
        if ext == ".docx":
            resume_paragraphs = read_docx_structured(resume_path)
        else:
            raw_text = read_pdf(resume_path)
            resume_paragraphs = [
                {"style": "Normal", "text": line}
                for line in raw_text.splitlines() if line.strip()
            ]
    except Exception as e:
        print(f"  [skip] Could not read source resume: {e}")
        return

    if not resume_paragraphs:
        print("  [skip] Source resume is empty.")
        return

    # Call Claude — full rewrite, 2-page target
    print("Calling Claude to rewrite resume (2-page target)...")
    try:
        new_paragraphs = call_claude_rewrite(new_jd_text, top_doc, resume_paragraphs)
    except Exception as e:
        print(f"  [skip] Claude error: {e}")
        return
    print(f"  {len(new_paragraphs)} paragraphs generated.")

    # Save output DOCX alongside the JD PDF in new_jobs/
    out_filename = f"tailored_{jd_stem}.docx"
    out_path = os.path.join(config.NEW_JD_DIR, out_filename)
    write_tailored_docx(new_paragraphs, out_path)
    print(f"  Saved → {out_path}")


def main():
    parser = argparse.ArgumentParser(description="Tailor resume for new job(s).")
    parser.add_argument("--jd", help="Path to a single JD PDF/TXT. Omit to process all PDFs in new_jobs/.")
    parser.add_argument("--top-k", type=int, default=3, help="Similar past JDs to retrieve (default 3).")
    args = parser.parse_args()

    if not config.ANTHROPIC_API_KEY:
        sys.exit("ANTHROPIC_API_KEY not set in .env")

    # Collect JD files to process
    if args.jd:
        jd_files = [args.jd]
    else:
        jd_files = collect_inbox_pdfs()
        if not jd_files:
            sys.exit(
                f"No PDF found in {config.NEW_JD_DIR}\n"
                "Drop job description PDF(s) there and re-run, or use --jd <path>."
            )
        print(f"Found {len(jd_files)} JD PDF(s) in new_jobs/:")
        for f in jd_files:
            print(f"  {os.path.basename(f)}")

    # Connect to ChromaDB once
    chroma = chromadb.PersistentClient(path=config.DB_PATH)
    collection = chroma.get_collection("job_descriptions")

    for jd_path in jd_files:
        print("\n" + "=" * 60)
        process_one_jd(jd_path, collection, args.top_k)

    print("\n" + "=" * 60)
    print(f"Done. Tailored DOCX(s) saved in: {config.NEW_JD_DIR}")
    print("=" * 60)


if __name__ == "__main__":
    main()
